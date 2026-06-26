"""
Post-process pandoc-generated docx:
  - adds thin gray borders to all tables
  - sets column widths per table (by position in document)

Usage:
    uv run python scripts/format_docx.py <file.docx>
    uv run python scripts/format_docx.py <input.docx> <output.docx>

The file is overwritten in-place if no output path is given.
"""

import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Text width on an A4 page with 2.5 cm margins on each side: 21 - 5 = 16 cm
# Column widths are expressed as fractions of this total.
PAGE_WIDTH_CM = 16.0

BORDER_COLOR = "BFBFBF"   # light gray
BORDER_SIZE  = 4           # in eighths of a point (4 = 0.5 pt)

# Column widths (fractions of PAGE_WIDTH_CM), indexed by table order in document.
# A row of fractions must sum to <= 1.0 (Pandoc/Word ignores the remainder).
TABLE_WIDTHS = [
    # 0 – 1.2 Arealtabellen: Areal | Kvm | Merknad
    [0.50, 0.12, 0.38],
    # 1 – 2.2 MVA-strategi: Segment | MVA-modell | Juridisk grunnlag
    [0.30, 0.25, 0.45],
    # 2 – 3.2 Konkurransekart: Aktør (lokasjon) | Pris/kvm/mnd | Differensiator
    [0.45, 0.18, 0.37],
    # 3 – 3.3 Målgrupper: Segment | Behov | FG30-tilbud
    [0.28, 0.42, 0.30],
    # 4 – 4.1 Lagerenheter: Kategori | Antall | Selvbetjent | Full-service
    [0.38, 0.13, 0.24, 0.25],
    # 5 – 4.3 Tilleggstjenester: Tjeneste | Pris | Volum est. | Rev. per år
    [0.38, 0.20, 0.20, 0.22],
    # 6 – 5.1 Belegg: Segment | År 1 | År 2 | År 3 | Stabilisert
    [0.32, 0.17, 0.17, 0.17, 0.17],
    # 7 – 5.2 Leieinntekter: Segment | År 1 | År 2 | År 3 | Stabilisert
    [0.28, 0.18, 0.18, 0.18, 0.18],
    # 8 – 6.1 Investeringsanslag: Post | Kostnad (MNOK)
    [0.80, 0.20],
    # 9 – 6.1 Finansiering: Kilde | Beløp (MNOK) | Merknad
    [0.40, 0.15, 0.45],
    # 10 – 6.2 Driftskostnader: Post | MNOK/år
    [0.80, 0.20],
    # 11 – 6.3 EBITDA-projeksjon: (label) | År 1 | År 2 | År 3 | Stab.
    [0.28, 0.18, 0.18, 0.18, 0.18],
    # 12 – 8.1 NPV: Scenario | Nettoinvestering | EBITDA stab. | Cap rate | Eiendomsverdi
    [0.22, 0.20, 0.20, 0.18, 0.20],
    # 13 – 9 Risikovurdering: Risiko | Sanns. | Konsekvens | Tiltak
    [0.40, 0.14, 0.18, 0.28],
    # 14 – 10 Fremdriftsplan: Milepæl | Dato (est.) | Status
    [0.55, 0.25, 0.20],
]


def _get_or_add(parent, tag):
    el = parent.find(qn(tag))
    if el is None:
        el = OxmlElement(tag)
        parent.append(el)
    return el


def _set_border_element(borders_el, edge):
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"),   "single")
    el.set(qn("w:sz"),    str(BORDER_SIZE))
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), BORDER_COLOR)
    borders_el.append(el)


def add_cell_borders(cell):
    tc    = cell._tc
    tcPr  = _get_or_add(tc, "w:tcPr")
    # Remove any existing tcBorders so we start clean
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        _set_border_element(tcBorders, edge)
    tcPr.append(tcBorders)


def set_column_widths(table, fractions):
    """Set column widths (in twips) for every cell in the table."""
    widths_twips = [int(f * PAGE_WIDTH_CM * 567) for f in fractions]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i >= len(widths_twips):
                break
            tc   = cell._tc
            tcPr = _get_or_add(tc, "w:tcPr")
            for old in tcPr.findall(qn("w:tcW")):
                tcPr.remove(old)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"),    str(widths_twips[i]))
            tcW.set(qn("w:type"), "dxa")
            # Insert before other properties so Word respects it
            tcPr.insert(0, tcW)


def process(input_path, output_path, borders_only=False):
    doc = Document(input_path)

    for i, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                add_cell_borders(cell)

        if not borders_only:
            if i < len(TABLE_WIDTHS):
                set_column_widths(table, TABLE_WIDTHS[i])
            else:
                print(f"  warning: no width config for table {i} – borders applied, widths left as-is")

    doc.save(output_path)
    print(f"Done: {len(doc.tables)} tables processed -> {output_path}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    args = [a for a in sys.argv[1:] if not a.startswith("--")]
    borders_only = "--borders-only" in sys.argv
    inp  = args[0]
    outp = args[1] if len(args) > 1 else inp
    process(inp, outp, borders_only=borders_only)
